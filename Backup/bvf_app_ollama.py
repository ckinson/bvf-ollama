# bvf_app_ollama_full.py
# Streamlit BVF Builder using local Ollama LLM
# Features:
#  - Strict JSON-only instructions
#  - Automatic JSON salvage if model adds extra text
#  - Curated, deduplicated output
#  - Model selector in UI
#  - Local PDF & DOCX uploads
#  - PDF export via reportlab

import os
import io
import json
from dataclasses import dataclass, field, asdict
from typing import List, Dict, Optional

import streamlit as st
import requests
from bs4 import BeautifulSoup
from readability import Document
from pdfminer.high_level import extract_text as pdf_extract_text
import pandas as pd
import plotly.graph_objects as go
from ollama import chat
from docx import Document as DocxDocument
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet

# ---------------------------
# Config
# ---------------------------
st.set_page_config(page_title="BVF Builder (Ollama)", layout="wide")

DEFAULT_INDUSTRY_STRATEGIES = [
    "Omni Channel Presence",
    "Speed to Customer",
    "Differentiated Customer Brand Experience",
    "Re-Define Store Operations",
    "Distribution & Logistics",
    "Adapt New Technologies"
]

# ---------------------------
# Data structures
# ---------------------------
@dataclass
class BVF:
    company: str
    executive_kpis: List[str] = field(default_factory=list)
    financial_operational_kpis: List[str] = field(default_factory=list)
    industry_strategies: Dict[str, List[str]] = field(default_factory=dict)
    business_processes_functions: Dict[str, List[str]] = field(default_factory=dict)
    operating_kpis: Dict[str, List[str]] = field(default_factory=dict)
    sources: List[str] = field(default_factory=list)

    def curate(self):
        """Deduplicate and shorten lists for better presentation."""
        self.executive_kpis = sorted(set(self.executive_kpis))[:8]
        self.financial_operational_kpis = sorted(set(self.financial_operational_kpis))[:10]
        self.industry_strategies = {k: sorted(set(v))[:8] for k, v in self.industry_strategies.items()}
        self.business_processes_functions = {k: sorted(set(v))[:8] for k, v in self.business_processes_functions.items()}
        self.operating_kpis = {k: sorted(set(v))[:8] for k, v in self.operating_kpis.items()}

    def to_frame(self) -> pd.DataFrame:
        rows = []
        for strat, inits in self.industry_strategies.items():
            rows.append({"Section": "Industry Strategy", "Lane": strat, "Items": "\n".join(inits)})
        for lane, items in self.business_processes_functions.items():
            rows.append({"Section": "Business Process & Functions", "Lane": lane, "Items": "\n".join(items)})
        for lane, items in self.operating_kpis.items():
            rows.append({"Section": "Operating KPIs", "Lane": lane, "Items": "\n".join(items)})
        return pd.DataFrame(rows)

# ---------------------------
# Helpers
# ---------------------------
def fetch_text(url: str) -> str:
    try:
        if url.lower().endswith(".pdf"):
            return pdf_extract_text(url)
        resp = requests.get(url, timeout=30, headers={"User-Agent": "Mozilla/5.0"})
        resp.raise_for_status()
        if "application/pdf" in resp.headers.get("Content-Type", ""):
            with open("temp.pdf", "wb") as fh:
                fh.write(resp.content)
            return pdf_extract_text("temp.pdf")
        doc = Document(resp.text)
        soup = BeautifulSoup(doc.summary(), "lxml")
        return soup.get_text(separator=" ", strip=True)
    except Exception:
        return ""

def read_uploaded_file(uploaded_file) -> str:
    if uploaded_file.name.lower().endswith(".pdf"):
        with open("temp_uploaded.pdf", "wb") as f:
            f.write(uploaded_file.read())
        return pdf_extract_text("temp_uploaded.pdf")
    elif uploaded_file.name.lower().endswith(".docx"):
        doc = DocxDocument(uploaded_file)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    else:
        return uploaded_file.read().decode("utf-8", errors="ignore")

def call_ollama(prompt: str, model: str) -> str:
    response = chat(model=model, messages=[{"role": "user", "content": prompt}])
    return response["message"]["content"]

def parse_json_safely(output: str) -> Optional[dict]:
    try:
        return json.loads(output)
    except json.JSONDecodeError:
        start = output.find("{")
        end = output.rfind("}")
        if start != -1 and end != -1:
            try:
                return json.loads(output[start:end+1])
            except json.JSONDecodeError:
                return None
        return None

def extract_bvf(company: str, corp_text: str, model: str) -> BVF:
    schema_hint = json.dumps({
        "executive_kpis": ["string"],
        "financial_operational_kpis": ["string"],
        "industry_strategies": {"Strategy Name": ["Initiative1", "Initiative2"]},
        "business_processes_functions": {"Lane Name": ["Project1", "Project2"]},
        "operating_kpis": {"Lane Name": ["KPI1", "KPI2"]},
        "sources": ["string"]
    }, indent=2)

    prompt = f"""
You are a senior business strategist.
Company: {company}

Task:
Based ONLY on the provided text, extract a **curated, concise, sector-specific** Business Value Framework (BVF):
- Use clear, business-relevant, and non-generic language.
- Remove duplicates.
- Prioritize most impactful KPIs, strategies, and initiatives.
- Keep each list short and to the point.

Return ONLY valid JSON exactly matching this structure:
{schema_hint}

Text to analyze:
{corp_text[:100000]}
"""

    output = call_ollama(prompt, model)
    data = parse_json_safely(output)

    if not data:
        st.error("❌ Model did not return valid JSON after salvage.")
        return BVF(company=company)

    bvf = BVF(company=company, **data)
    bvf.curate()
    return bvf

def render_bvf_figure(bvf: BVF) -> go.Figure:
    strategies = list(bvf.industry_strategies.keys())
    n_cols = max(5, len(strategies))
    n_rows = 6

    fig = go.Figure()
    fig.update_xaxes(visible=False, range=[0, n_cols])
    fig.update_yaxes(visible=False, range=[0, n_rows])
    fig.update_layout(height=900, margin=dict(l=40, r=40, t=60, b=40), showlegend=False)

    def box(x0, y0, x1, y1, title, body):
        fig.add_shape(type="rect", x0=x0, y0=y0, x1=x1, y1=y1, line=dict(width=1))
        fig.add_annotation(x=(x0 + x1) / 2, y=y1 - 0.3, text=f"<b>{title}</b>", showarrow=False, yanchor="top")
        fig.add_annotation(x=(x0 + x1) / 2, y=y1 - 0.6, text=body.replace("\n", "<br>"), showarrow=False, yanchor="top")

    y = n_rows
    fig.add_annotation(x=n_cols / 2, y=y - 0.5, text=f"<b>Business Value Framework — {bvf.company}</b>",
                       showarrow=False, yanchor="top", font=dict(size=20))
    y -= 1

    box(0, y - 1, n_cols, y, "Executive KPIs", "\n".join(bvf.executive_kpis))
    y -= 1
    box(0, y - 1, n_cols, y, "Financial / Operational KPIs", "\n".join(bvf.financial_operational_kpis))
    y -= 1

    col_w = n_cols / len(strategies) if strategies else n_cols
    for i, strat in enumerate(strategies):
        x0 = i * col_w
        x1 = (i + 1) * col_w
        items = bvf.industry_strategies.get(strat, [])
        box(x0, y - 1, x1, y, strat, "\n".join(items))
    y -= 1

    lanes = list(bvf.business_processes_functions.keys()) or ["Processes"]
    col_w = n_cols / len(lanes)
    for i, lane in enumerate(lanes):
        x0 = i * col_w
        x1 = (i + 1) * col_w
        items = bvf.business_processes_functions.get(lane, [])
        box(x0, y - 1, x1, y, lane, "\n".join(items))
    y -= 1

    lanes2 = list(bvf.operating_kpis.keys()) or lanes
    col_w = n_cols / len(lanes2)
    for i, lane in enumerate(lanes2):
        x0 = i * col_w
        x1 = (i + 1) * col_w
        items = bvf.operating_kpis.get(lane, [])
        box(x0, y - 1, x1, y, f"{lane} — Operating KPIs", "\n".join(items))
    y -= 1

    return fig

def export_bvf_to_pdf(bvf: BVF, filename: str):
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(filename, pagesize=A4)
    story = []

    story.append(Paragraph(f"Business Value Framework — {bvf.company}", styles["Title"]))
    story.append(Spacer(1, 12))

    def add_section(title, items):
        story.append(Paragraph(f"<b>{title}</b>", styles["Heading2"]))
        if isinstance(items, list):
            for i in items:
                story.append(Paragraph(f"- {i}", styles["Normal"]))
        elif isinstance(items, dict):
            for k, v in items.items():
                story.append(Paragraph(f"<b>{k}</b>", styles["Heading3"]))
                for i in v:
                    story.append(Paragraph(f"- {i}", styles["Normal"]))
        story.append(Spacer(1, 12))

    add_section("Executive KPIs", bvf.executive_kpis)
    add_section("Financial / Operational KPIs", bvf.financial_operational_kpis)
    add_section("Industry Strategies", bvf.industry_strategies)
    add_section("Business Processes & Functions", bvf.business_processes_functions)
    add_section("Operating KPIs", bvf.operating_kpis)

    doc.build(story)

# ---------------------------
# UI
# ---------------------------
st.title("🧭 BVF Builder ")

company = st.text_input("Company name", placeholder="e.g., Aviva Insurance")
model_name = st.selectbox("Ollama model to use", ["llama3", "mistral", "gemma", "qwen"], index=0)

manual_urls = st.text_area("Optional: paste specific URLs (one per line)", height=100)
manual_text = st.text_area("Paste raw strategy text here", height=200)

uploaded_files = st.file_uploader("Upload local PDF or DOCX strategy files", type=["pdf", "docx"], accept_multiple_files=True)

if st.button("Load Uploaded Files"):
    file_text = ""
    for file in uploaded_files:
        file_text += "\n\nSOURCE: " + file.name + "\n" + read_uploaded_file(file)
    manual_text += "\n" + file_text

if st.button("Fetch from URLs"):
    all_text = ""
    for u in manual_urls.splitlines():
        if u.strip():
            all_text += f"\n\nSOURCE: {u}\n{fetch_text(u.strip())}"
    manual_text += "\n" + all_text

if st.button("Build BVF (Local Ollama)", disabled=not company or not manual_text.strip()):
    with st.spinner("Generating BVF using local Ollama..."):
        bvf = extract_bvf(company, manual_text, model=model_name)
        st.session_state["bvf"] = bvf
        if bvf.executive_kpis:
            st.success("✅ BVF generated!")
            st.json(asdict(bvf))

bvf: Optional[BVF] = st.session_state.get("bvf")
if bvf and bvf.executive_kpis:
    fig = render_bvf_figure(bvf)
    st.plotly_chart(fig, use_container_width=True)
    df = bvf.to_frame()
    st.download_button("Download JSON", json.dumps(asdict(bvf), indent=2), file_name=f"{bvf.company}_BVF.json")
    st.download_button("Download CSV", df.to_csv(index=False).encode("utf-8"), file_name=f"{bvf.company}_BVF.csv", mime="text/csv")

    # PDF Export
    pdf_filename = f"{bvf.company}_BVF.pdf"
    export_bvf_to_pdf(bvf, pdf_filename)
    with open(pdf_filename, "rb") as pdf_file:
        st.download_button("Download PDF", pdf_file, file_name=pdf_filename, mime="application/pdf")
